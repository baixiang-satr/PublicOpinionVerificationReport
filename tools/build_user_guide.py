"""Build the portable release user guide as a polished DOCX."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from word_list_restart import apply_num_id, new_number_list

OUTPUT = Path(r"D:\舆情验证报告工具用户使用说明书.docx")
FONT = "Calibri"
CN_FONT = "Microsoft YaHei"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
BODY = "222222"
MUTED = "666666"
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "FFF4CE"
TABLE_BORDER = "B8C3D1"
WIDTH_DXA = 9360
INDENT_DXA = 120


def set_run(run, *, size=11, color=BODY, bold=False, italic=False, font=FONT):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CN_FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    return run


def configure_style(style, *, size, color, before, after, line=1.25, bold=False):
    style.font.name = FONT
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CN_FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_style(doc.styles["Normal"], size=11, color=BODY, before=0, after=6)
    configure_style(doc.styles["Heading 1"], size=16, color=BLUE, before=18, after=10, bold=True)
    configure_style(doc.styles["Heading 2"], size=13, color=BLUE, before=14, after=7, bold=True)
    configure_style(doc.styles["Heading 3"], size=12, color=DARK_BLUE, before=10, after=5, bold=True)
    for name in ("List Bullet", "List Number"):
        configure_style(doc.styles[name], size=11, color=BODY, before=0, after=4)
        doc.styles[name].paragraph_format.left_indent = Inches(0.375)
        doc.styles[name].paragraph_format.first_line_indent = Inches(-0.188)
    doc.core_properties.title = "舆情验证报告工具用户使用说明书"
    doc.core_properties.subject = "Windows 便携版安装、登录、采集、补录与导出说明"
    doc.core_properties.author = "舆情验证报告工具项目组"


def add_list(doc, text: str, num_id: int, *, numbered=False):
    p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    if numbered:
        apply_num_id(p, num_id)
    set_run(p.add_run(text))
    return p


def add_body(doc, text: str, *, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        set_run(p.add_run(bold_prefix), bold=True)
        set_run(p.add_run(text[len(bold_prefix) :]))
    else:
        set_run(p.add_run(text))
    return p


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    width = tbl_pr.first_child_found_in("w:tblW")
    width.set(qn("w:w"), str(WIDTH_DXA))
    width.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), str(INDENT_DXA))
    indent.set(qn("w:type"), "dxa")
    tbl_pr.append(indent)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        item = OxmlElement(f"w:{edge}")
        item.set(qn("w:val"), "single")
        item.set(qn("w:sz"), "4")
        item.set(qn("w:color"), TABLE_BORDER)
        borders.append(item)
    tbl_pr.append(borders)
    grid = table._tbl.tblGrid
    for item, item_width in zip(grid.gridCol_lst, widths, strict=True):
        item.set(qn("w:w"), str(item_width))
    for row in table.rows:
        for cell, item_width in zip(row.cells, widths, strict=True):
            cell.width = Inches(item_width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(item_width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    header_row = table.rows[0]
    tr_pr = header_row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)
    for cell, text in zip(header_row.cells, headers, strict=True):
        cell._tc.get_or_add_tcPr().append(_fill(HEADER_FILL))
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(text), size=9.5, bold=True, color=DARK_BLUE)
    for values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, values, strict=True):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            set_run(p.add_run(text), size=9.5)
    set_table_geometry(table, widths)
    return table


def _fill(color: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    return shading


def add_callout(doc, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [WIDTH_DXA])
    cant_split = OxmlElement("w:cantSplit")
    table.rows[0]._tr.get_or_add_trPr().append(cant_split)
    cell = table.cell(0, 0)
    cell._tc.get_or_add_tcPr().append(_fill(CALLOUT_FILL))
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_together = True
    set_run(p.add_run(f"{title}  "), bold=True, color=DARK_BLUE)
    set_run(p.add_run(text), size=10.5)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))
    set_run(run, size=8.5, color=MUTED)


def set_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("舆情验证报告工具｜用户使用说明书"), size=8.5, color=MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    set_run(fp.add_run("第 "), size=8.5, color=MUTED)
    add_field(fp, "PAGE")
    set_run(fp.add_run(" 页"), size=8.5, color=MUTED)


def add_cover(doc: Document) -> None:
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(kicker.add_run("WINDOWS 便携版｜操作手册"), size=10.5, color=BLUE, bold=True)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(10)
    set_run(title.add_run("舆情验证报告工具"), size=28, color=DARK_BLUE, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(34)
    set_run(subtitle.add_run("用户使用说明书"), size=16, color=BLUE)
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(desc.add_run("批量导入 URL｜登录态复用｜证据截图｜人工补录｜固定模板导出"), size=11, color=MUTED)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(70)
    set_run(meta.add_run(f"版本 1.0　｜　发布日期 {date(2026, 8, 1).isoformat()}"), size=10, color=MUTED)
    doc.add_page_break()


def build() -> Path:
    doc = Document()
    configure_document(doc)
    set_header_footer(doc)
    bullet_id = 0
    number_id = new_number_list(doc)
    add_cover(doc)

    doc.add_heading("1. 一分钟快速开始", level=1)
    for text in (
        "完整解压便携包，双击“舆情验证报告工具.exe”。不要只复制 EXE。",
        "准备包含 http(s) 链接的 TXT、CSV 或 XLSX 文件。",
        "进入“管理平台登录态”，把本次涉及的平台逐个更新到“登录态有效”。",
        "选择输入文件，保持默认参数并开始抓取。",
        "在结果页检查待补录项，必要时编辑字段或补充内容页/个人页截图。",
        "在“预览导出”生成 template.zip；该 ZIP 才是正式交付包。",
    ):
        add_list(doc, text, number_id, numbered=True)
    add_callout(doc, "最重要的原则", "工具不会把游客页当作合格证据。平台显示“需要重新登录”时，先更新登录态，再开始采集。")

    doc.add_heading("2. 安装、迁移与运行环境", level=1)
    add_table(
        doc,
        ["项目", "要求"],
        [
            ["操作系统", "Windows 10 / 11，64 位"],
            ["内存", "至少 4 GB，建议 8 GB 或以上"],
            ["网络", "可访问目标平台；公司代理或防火墙可能影响加载"],
            ["依赖", "无需安装 Python、Node.js 或浏览器；组件已随包提供"],
            ["界面运行时", "系统需有 Microsoft Edge WebView2 Runtime（Win10/11 通常自带）"],
        ],
        [1700, 7660],
    )
    add_body(doc, "把压缩包解压到本地普通目录（例如 D:\\舆情工具），确保当前用户有写入权限。不要在压缩包内直接运行，也不要删除 template、web、ms-playwright 或 poir_ocr_worker.exe。")
    add_callout(doc, "换电脑须知", "登录态使用 Windows 当前用户的 DPAPI 加密，不能跨电脑或跨 Windows 用户复制。迁移便携包后，各平台必须在新电脑重新登录一次。")
    for text in (
        "首次启动约需 5–15 秒。若 SmartScreen 提示，请核对文件来源和发布包哈希后选择运行；不要关闭系统安全防护。",
        "若窗口空白，安装或修复 Microsoft Edge WebView2 Runtime 后重试。",
        "运行过程中不要移动程序目录；任务结束后再复制或归档输出。",
    ):
        add_list(doc, text, bullet_id)

    doc.add_heading("3. 准备 URL 输入文件", level=1)
    add_body(doc, "程序会从文件中提取所有 http:// 或 https:// 链接，并自动去重。每条 URL 对应一条证据记录。")
    add_table(
        doc,
        ["格式", "建议写法", "说明"],
        [
            ["TXT", "每行一个 URL", "最稳定，适合日常批量任务"],
            ["CSV", "单独一列保存 URL", "允许有表头；非 URL 内容会被忽略"],
            ["XLSX", "任意工作表的单元格内放 URL", "用于已有清单；合并单元格应避免"],
        ],
        [1200, 3000, 5160],
    )
    for text in (
        "尽量使用可在桌面浏览器打开的原始内容链接，不要使用只能在手机 App 内部识别的截图或口令。",
        "同一平台可放多条 URL；程序会按平台串行复用同一个已登录 context，其他平台可并行。",
        "删除过期、无权限或已知失效链接，可显著减少待补录项。",
    ):
        add_list(doc, text, bullet_id)

    doc.add_heading("4. 管理平台登录态", level=1)
    add_body(doc, "这是正式采集前的必做步骤。工具按平台隔离保存登录态，并在抓取前用新的浏览器 context 复验。")
    number_id = new_number_list(doc)
    for text in (
        "点击“管理平台登录态…”，优先处理“本次 URL 涉及的平台”。",
        "在目标平台卡片中点击“登录 / 更新”。浏览器先在屏幕外加载，页面稳定后只显示一次。",
        "在官方页面完成扫码、验证码或密码登录；工具不会自动提交验证码。",
        "回到登录态管理中心，点击“完成登录并保存”。",
        "等待卡片显示“登录态有效”，再处理下一个平台。",
    ):
        add_list(doc, text, number_id, numbered=True)
    add_callout(doc, "微信视频号", "必须在视频号官方登录页完成二维码登录，并让保存态包含账号级 sessionid。仅打开公开分享页或只有统计缓存，不代表已登录；这种旧状态会显示“需要重新登录”。")
    add_callout(doc, "隐私保护", "Cookie/Token 按平台加密保存在当前 Windows 用户的 LocalAppData，不写入日志，也不会进入 template.zip。")

    doc.add_heading("5. 运行参数", level=1)
    add_table(
        doc,
        ["参数", "默认/建议", "作用"],
        [
            ["同时处理", "3（建议 3–5）", "跨平台并行上限；过大更容易触发风控"],
            ["单页超时", "45 秒", "页面在规定时间内未就绪则进入重试/待补录"],
            ["失败重试", "1 次", "导航或截图等瞬时失败会自动再尝试"],
            ["截图格式", "JPG", "体积较小；文字精细时可选 PNG"],
            ["浏览器模式", "后台有界面（固定）", "保持真实浏览器指纹并复用登录态"],
        ],
        [1600, 1900, 5860],
    )
    add_body(doc, "一般不要为了追求速度把并发调高。批量质量下降时，先把“同时处理”降到 2–3，并保留 1 次重试。")

    doc.add_heading("6. 开始抓取与进度判断", level=1)
    for text in (
        "点击“开始抓取”后保持应用和网络在线，不要关闭应用或强制结束浏览器进程。",
        "程序会依次进行登录态预检、页面导航、字段提取、OCR、内容截图、作者主页截图和模板写入。",
        "同平台某条记录确认遇到登录屏障后，该平台剩余 URL 会暂停；其他平台继续。",
        "点击取消会安全停止后续任务；已完成记录会保留在检查点中。",
    ):
        add_list(doc, text, bullet_id)
    add_callout(doc, "如何理解“完成”", "“任务完成”不等于所有 URL 都成功。请在结果页查看“成功/待补录/失败”数量，并打开红色或黄色提示的记录。")

    doc.add_heading("7. 结果检查与人工补录", level=1)
    add_body(doc, "抓取结束后先预览表格，再进入“采集与补录”。红色空格通常表示模板必填字段仍缺失。")
    for text in (
        "核对标题、昵称/账号、发布时间、正文和发布平台是否与原页一致。",
        "点击蓝色 URL 可打开原页面复核；只在有权访问和处理的范围内补录。",
        "可编辑字段保存后会立即进入当前任务副本，不会修改只读基准模板。",
        "使用“重试失败项”重新处理瞬时超时；永久删除、无权限或仅 App 可见页面应人工确认。",
    ):
        add_list(doc, text, bullet_id)

    doc.add_heading("8. 截图规则：自动与人工", level=1)
    add_body(doc, "自动截图优先保留桌面登录页面的关键证据区；长文章会生成 1440 像素宽的长图。作者主页只有在身份匹配后才会保存。")
    add_body(doc, "人工框选步骤：", bold_prefix="人工框选步骤：")
    number_id = new_number_list(doc)
    for text in (
        "选中要补录的记录，点击“截取内容页”或“截取个人页”。",
        "在打开的已登录浏览器中滚动到目标位置，确认昵称、账号、时间或正文已完整可见。",
        "点击“开始框选”；此时画面冻结，拖动框选需要保留的区域。",
        "确认边界没有越出冻结画面，再保存；系统会自动裁剪到有效图像范围并关联文件名。",
    ):
        add_list(doc, text, number_id, numbered=True)
    add_callout(doc, "截图验收", "截图必须能看出内容主体和必要身份信息，不能只截空白、登录框、游客提示或被固定导航遮住的区域。发现偏移时取消本次框选，重新滚动定位后再截。")

    doc.add_heading("9. 导出与文件结构", level=1)
    add_body(doc, "在“预览导出”确认无误后生成 template.zip。ZIP 内固定包含 template/ 顶层目录、template.xlsx 以及被工作簿引用的截图/附件。")
    add_table(
        doc,
        ["文件/目录", "用途"],
        [
            ["output/<任务编号>/template.zip", "正式交付包"],
            ["quality_report.md", "本次质量明细和错误代码"],
            ["quality_summary.json", "可机器读取的覆盖率统计"],
            ["pending_manual_entry.csv", "需要人工补录的记录清单"],
            ["job_checkpoint.json", "断点、重试和审计信息"],
        ],
        [3000, 6360],
    )
    for text in (
        "不要手工改 ZIP 内部目录层级，也不要把 Cookie、浏览器 profile 或日志放入交付包。",
        "若需继续补录，在欢迎页选择“上传 template.zip 补录”，编辑后重新导出。",
        "交付前随机打开数条截图，并核对 Excel 内文件名均能在 ZIP 中找到。",
    ):
        add_list(doc, text, bullet_id)

    doc.add_heading("10. 状态与常见错误", level=1)
    add_table(
        doc,
        ["状态/提示", "含义", "处理"],
        [
            ["登录态有效", "已在新 context 中复验", "可以开始该平台采集"],
            ["需要重新登录", "缺少账号 Cookie 或状态失效", "登录 / 更新后确认保存"],
            ["待补录", "至少一项必填证据缺失", "复核登录态、重试或人工补录"],
            ["页面截图失败", "页面未稳定或瞬时异常", "使用重试；仍失败则人工截图"],
            ["作者身份不匹配", "主页账号与内容作者不一致", "不要强行保存，人工核对作者主页"],
            ["访问验证/风控", "平台要求验证码或限制访问", "降低并发，稍后在官方页面人工处理"],
        ],
        [1900, 3100, 4360],
    )

    doc.add_heading("11. 故障排查", level=1)
    for title, detail in (
        ("双击无反应", "等待 15 秒；检查文件是否完整解压、是否被安全软件隔离、WebView2 是否可用。"),
        ("登录窗口闪一下", "确认使用的是本说明书对应新版；关闭残留浏览器进程后重试。新版只会在页面完成加载后显示同一个窗口。"),
        ("批量比单条少数据", "把并发降到 2–3，保留 1 次重试；查看 quality_report.md，优先重试截图/OCR超时记录。"),
        ("截图偏移或裁切", "自动截图请重试；人工截图应在冻结画面中重新框选，不能把选区拖出画面边界。"),
        ("仍显示游客", "回到登录态中心更新该平台，确认“登录态有效”。不要在公开分享页停留后直接确认。"),
        ("微信视频号无法采集", "二维码登录视频号官方页并保存；若仍缺 sessionid，退出登录后重新扫码。"),
        ("OCR 没有识别", "确认 poir_ocr_worker.exe 与主程序同目录；图片过小、模糊或无文字时会跳过。"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        set_run(p.add_run(f"{title}："), bold=True, color=DARK_BLUE)
        set_run(p.add_run(detail))

    doc.add_heading("12. 合规、安全与交付前检查", level=1)
    for text in (
        "仅采集你有权访问、保存和处理的信息；不要绕过验证码、付费墙或站点权限。",
        "不要把程序的 LocalAppData 登录态目录、Cookie、Token 或调试文件交给他人。",
        "登录态疑似泄露时，在目标平台退出全部设备，并在工具中点击“退出登录”删除本机保存态。",
        "新电脑首次使用时逐平台登录；确认系统时间和时区正确，否则发布时间可能显示异常。",
        "交付前确认成功率/字段覆盖率、截图画面、待补录清单和 template.zip 结构。",
    ):
        add_list(doc, text, bullet_id)
    add_callout(doc, "完成标准", "本次任务中允许存在确有外部原因的待补录记录，但应能明确说明原因；其余记录的必填字段和内容页截图应完整，且不得用游客截图冒充登录证据。")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
